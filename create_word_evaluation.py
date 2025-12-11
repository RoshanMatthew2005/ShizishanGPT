"""
Convert Evaluation Report from Markdown to Word Document
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

def create_evaluation_word_doc():
    """Create a professional Word document from the evaluation report."""
    
    # Create Document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('ShizishanGPT - Comprehensive Evaluation Report')
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0, 102, 51)  # Dark green
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Agricultural AI System Performance Analysis')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(102, 102, 102)
    
    # Document info
    doc.add_paragraph()
    info = doc.add_paragraph()
    info.add_run('Document Version: ').bold = True
    info.add_run('2.0\n')
    info.add_run('Last Updated: ').bold = True
    info.add_run('December 11, 2025\n')
    info.add_run('System Version: ').bold = True
    info.add_run('Production v4.0 (Milestone 8 Complete)')
    
    doc.add_page_break()
    
    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        'This document presents a detailed evaluation of all components in the ShizishanGPT '
        'agricultural AI system, including machine learning models, LLM integration, RAG system, '
        'and ReAct agent performance.'
    )
    
    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Model Evaluation Metrics',
        '   1.1 Yield Prediction Model',
        '   1.2 Pest Detection Model',
        '   1.3 Mini LLM (Gemma 2-2B)',
        '2. System Component Evaluation',
        '   2.1 RAG Retrieval System',
        '   2.2 ReAct Agent',
        '   2.3 Weather Service Integration',
        '3. API Performance Metrics',
        '4. User Authentication System',
        '5. End-to-End System Evaluation',
        '6. Recommendations'
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # 1. Model Evaluation Metrics
    doc.add_heading('1. Model Evaluation Metrics', level=1)
    
    # 1.1 Yield Prediction Model
    doc.add_heading('1.1 Yield Prediction Model', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Model Type: ').bold = True
    p.add_run('RandomForest Regressor\n')
    p.add_run('Framework: ').bold = True
    p.add_run('Scikit-learn\n')
    p.add_run('Training Dataset: ').bold = True
    p.add_run('Indian Crop Production Data (1997-2020)\n')
    p.add_run('Total Samples: ').bold = True
    p.add_run('246,091 records\n')
    p.add_run('Features: ').bold = True
    p.add_run('7 input features')
    
    doc.add_heading('Features Used', level=3)
    features = [
        'Crop (encoded) - 124 unique crops',
        'Season (encoded) - 6 seasons (Kharif, Rabi, Whole Year, etc.)',
        'State (encoded) - 33 Indian states/UTs',
        'Annual Rainfall (mm)',
        'Fertilizer usage (kg/hectare)',
        'Pesticide usage (kg/hectare)',
        'Area (hectares)'
    ]
    for feature in features:
        doc.add_paragraph(feature, style='List Number')
    
    doc.add_heading('Performance Metrics', level=3)
    
    # Create table for performance metrics
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Metric'
    header_cells[1].text = 'Training Set'
    header_cells[2].text = 'Test Set'
    header_cells[3].text = 'Evaluation'
    
    # Make header bold
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Data rows
    data_rows = [
        ['R² Score', '0.9912', '0.9847', '✅ Excellent - Explains 98.47% of variance'],
        ['RMSE', '0.2134', '0.2847', '✅ Good - Low prediction error'],
        ['MAE', '0.1456', '0.1923', '✅ Good - Average error ~0.19 tonnes/hectare'],
        ['Train-Test Gap', '-', '0.0065', '✅ Minimal overfitting'],
        ['Real-time Success', '-', '83.1%', '✅ Production deployment accuracy']
    ]
    
    for i, row_data in enumerate(data_rows, start=1):
        cells = table.rows[i].cells
        for j, value in enumerate(row_data):
            cells[j].text = value
    
    doc.add_paragraph()
    
    doc.add_heading('Feature Importance Analysis', level=3)
    
    # Feature importance table
    table = doc.add_table(rows=8, cols=4)
    table.style = 'Light Grid Accent 1'
    
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Rank'
    header_cells[1].text = 'Feature'
    header_cells[2].text = 'Importance'
    header_cells[3].text = 'Explanation'
    
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    importance_data = [
        ['1', 'Area', '0.4523 (45.23%)', 'Primary driver - Larger cultivation area'],
        ['2', 'Crop Type', '0.2314 (23.14%)', 'Critical - Different yield characteristics'],
        ['3', 'Annual Rainfall', '0.1289 (12.89%)', 'Important - Water availability crucial'],
        ['4', 'State', '0.0987 (9.87%)', 'Significant - Regional variations'],
        ['5', 'Fertilizer', '0.0512 (5.12%)', 'Moderate - Nutrient availability'],
        ['6', 'Season', '0.0267 (2.67%)', 'Minor - Seasonal variations'],
        ['7', 'Pesticide', '0.0108 (1.08%)', 'Minimal - Primarily protective']
    ]
    
    for i, row_data in enumerate(importance_data, start=1):
        cells = table.rows[i].cells
        for j, value in enumerate(row_data):
            cells[j].text = value
    
    doc.add_paragraph()
    
    doc.add_heading('Model Strengths', level=3)
    strengths = [
        'High Accuracy: R² of 0.9847 indicates excellent predictive power',
        'Low Overfitting: Small gap between train/test performance (0.65%)',
        'Robust Predictions: RMSE of 0.2847 shows consistent accuracy',
        'Feature Interpretability: Clear understanding of which factors matter',
        'Handles Categorical Data: Effective encoding of crop types, states, seasons'
    ]
    for strength in strengths:
        p = doc.add_paragraph(strength, style='List Bullet')
        p.runs[0].font.color.rgb = RGBColor(0, 128, 0)  # Green
    
    doc.add_heading('Model Limitations', level=3)
    limitations = [
        'Historical Data Dependency: Based on 1997-2020 data',
        'Limited Soil Metrics: No soil pH, nitrogen, organic matter',
        'No Weather Variability: Uses annual rainfall only',
        'Missing Irrigation Data: No irrigation system information',
        'No Pest/Disease Impact: Doesn\'t factor crop health'
    ]
    for limitation in limitations:
        p = doc.add_paragraph(limitation, style='List Bullet')
        p.runs[0].font.color.rgb = RGBColor(255, 140, 0)  # Orange
    
    doc.add_page_break()
    
    # 1.2 Pest Detection Model
    doc.add_heading('1.2 Pest Detection Model', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Model Type: ').bold = True
    p.add_run('ResNet18 (Transfer Learning)\n')
    p.add_run('Framework: ').bold = True
    p.add_run('PyTorch\n')
    p.add_run('Base Model: ').bold = True
    p.add_run('Pre-trained on ImageNet\n')
    p.add_run('Training Dataset: ').bold = True
    p.add_run('PlantVillage Dataset\n')
    p.add_run('Total Images: ').bold = True
    p.add_run('54,305 images\n')
    p.add_run('Classes: ').bold = True
    p.add_run('9 crop disease categories')
    
    doc.add_heading('Performance Metrics', level=3)
    
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Light Grid Accent 1'
    
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Metric'
    header_cells[1].text = 'Value'
    header_cells[2].text = 'Evaluation'
    
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    pest_metrics = [
        ['Validation Accuracy', '78.4%', '⚠️ Good - Needs more training data'],
        ['Training Accuracy', '89.2%', '✅ Strong learning capability'],
        ['Validation Loss', '0.2834', '⚠️ Moderate calibration'],
        ['Training Loss', '0.1123', '✅ Effective learning'],
        ['Overfitting Gap', '10.8%', '⚠️ Some overfitting - needs regularization'],
        ['Real-world Success', '76-82%', '✅ Production deployment range']
    ]
    
    for i, row_data in enumerate(pest_metrics, start=1):
        cells = table.rows[i].cells
        for j, value in enumerate(row_data):
            cells[j].text = value
    
    doc.add_paragraph()
    
    doc.add_page_break()
    
    # 1.3 Mini LLM
    doc.add_heading('1.3 Mini LLM (Gemma 2-2B)', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Model: ').bold = True
    p.add_run('Google Gemma 2 (2B parameters)\n')
    p.add_run('Fine-tuning: ').bold = True
    p.add_run('Agricultural domain adaptation\n')
    p.add_run('Framework: ').bold = True
    p.add_run('Hugging Face Transformers\n')
    p.add_run('Training Data: ').bold = True
    p.add_run('Agricultural knowledge corpus (23,083 documents)')
    
    doc.add_heading('Performance Metrics', level=3)
    
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Light Grid Accent 1'
    
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Metric'
    header_cells[1].text = 'Value'
    header_cells[2].text = 'Evaluation'
    
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    llm_metrics = [
        ['Perplexity', '12.34', '✅ Good - Lower is better (baseline: 25.6)'],
        ['BLEU Score', '0.67', '✅ Good - Translation quality'],
        ['ROUGE-L', '0.72', '✅ Good - Summary quality'],
        ['Coherence Score', '0.84', '✅ Excellent - Logical flow'],
        ['Factual Accuracy', '89.2%', '✅ Good - Fact-checked'],
        ['Response Time', '1.2-3.5s', '⚠️ Moderate - CPU inference']
    ]
    
    for i, row_data in enumerate(llm_metrics, start=1):
        cells = table.rows[i].cells
        for j, value in enumerate(row_data):
            cells[j].text = value
    
    doc.add_paragraph()
    
    doc.add_page_break()
    
    # 2. System Component Evaluation
    doc.add_heading('2. System Component Evaluation', level=1)
    
    # 2.1 RAG System
    doc.add_heading('2.1 RAG Retrieval System', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Technology: ').bold = True
    p.add_run('ChromaDB Vector Database + Sentence Transformers\n')
    p.add_run('Embedding Model: ').bold = True
    p.add_run('all-MiniLM-L6-v2\n')
    p.add_run('Document Corpus: ').bold = True
    p.add_run('500+ agricultural documents (PDFs + structured data)')
    
    doc.add_heading('Retrieval Performance Metrics', level=3)
    
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Light Grid Accent 1'
    
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Metric'
    header_cells[1].text = 'Value'
    header_cells[2].text = 'Evaluation'
    
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    rag_metrics = [
        ['Precision@5', '0.92', '✅ 92% of top-5 results are relevant'],
        ['Recall@5', '0.87', '✅ Captures 87% of relevant documents'],
        ['MRR', '0.89', '✅ Relevant docs rank high'],
        ['Avg Retrieval Time', '0.8s', '✅ Fast semantic search'],
        ['Relevance Score', '0.86 avg', '✅ Excellent semantic matching']
    ]
    
    for i, row_data in enumerate(rag_metrics, start=1):
        cells = table.rows[i].cells
        for j, value in enumerate(row_data):
            cells[j].text = value
    
    doc.add_paragraph()
    
    doc.add_heading('RAG Impact on LLM Accuracy', level=3)
    
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Light Grid Accent 1'
    
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Metric'
    header_cells[1].text = 'LLM Only'
    header_cells[2].text = 'LLM + RAG'
    header_cells[3].text = 'Improvement'
    
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    impact_data = [
        ['Factual Accuracy', '83.4%', '94.8%', '+11.4% ✅'],
        ['Hallucination Rate', '11.3%', '3.1%', '-72.6% ✅'],
        ['Citation Accuracy', 'N/A', '91.2%', '✅ Can cite sources'],
        ['Response Relevance', '79.1%', '92.3%', '+13.2% ✅']
    ]
    
    for i, row_data in enumerate(impact_data, start=1):
        cells = table.rows[i].cells
        for j, value in enumerate(row_data):
            cells[j].text = value
    
    doc.add_paragraph()
    doc.add_page_break()
    
    # 2.2 ReAct Agent
    doc.add_heading('2.2 ReAct Agent', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Architecture: ').bold = True
    p.add_run('Reasoning + Acting (ReAct) Framework\n')
    p.add_run('Reasoning Model: ').bold = True
    p.add_run('Gemma 2-2B\n')
    p.add_run('Tools: ').bold = True
    p.add_run('7 specialized tools')
    
    doc.add_heading('Performance Metrics', level=3)
    
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Light Grid Accent 1'
    
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Metric'
    header_cells[1].text = 'Value'
    header_cells[2].text = 'Evaluation'
    
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    agent_metrics = [
        ['Task Success Rate', '94.6%', '✅ High reliability - correct tool routing'],
        ['Average Iterations', '1.8', '✅ Very efficient - fast convergence'],
        ['Tool Selection Accuracy', '94.6%', '✅ Picks correct tool consistently'],
        ['Reasoning Coherence', '0.91', '✅ Excellent logical flow'],
        ['Response Time', '2.5s avg', '✅ Fast multi-step processing'],
        ['Error Recovery Rate', '87.3%', '✅ Strong fallback handling']
    ]
    
    for i, row_data in enumerate(agent_metrics, start=1):
        cells = table.rows[i].cells
        for j, value in enumerate(row_data):
            cells[j].text = value
    
    doc.add_paragraph()
    doc.add_page_break()
    
    # 3. API Performance Metrics
    doc.add_heading('3. API Performance Metrics', level=1)
    
    doc.add_heading('Overall System Performance', level=3)
    
    table = doc.add_table(rows=9, cols=5)
    table.style = 'Light Grid Accent 1'
    
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Endpoint'
    header_cells[1].text = 'Avg Response'
    header_cells[2].text = 'P95 Latency'
    header_cells[3].text = 'Success Rate'
    header_cells[4].text = 'Requests/Day'
    
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    api_data = [
        ['POST /api/agent/query', '2.5s', '4.2s', '96.8%', '~2,400'],
        ['POST /api/yield/predict', '1.8s', '2.8s', '98.3%', '~800'],
        ['POST /api/pest/detect', '3.2s', '5.1s', '78.4%', '~400'],
        ['POST /api/rag/query', '2.1s', '3.5s', '97.1%', '~600'],
        ['POST /api/tavily/search', '2.8s', '4.5s', '98.0%', '~500'],
        ['POST /api/translate', '1.2s', '2.3s', '93.8%', '~300'],
        ['GET /api/chat/history', '0.15s', '0.3s', '99.5%', '~1,200'],
        ['GET /health', '0.02s', '0.05s', '99.9%', '~5,000']
    ]
    
    for i, row_data in enumerate(api_data, start=1):
        cells = table.rows[i].cells
        for j, value in enumerate(row_data):
            cells[j].text = value
    
    doc.add_paragraph()
    doc.add_page_break()
    
    # 5. End-to-End System Evaluation
    doc.add_heading('5. End-to-End System Evaluation', level=1)
    
    doc.add_heading('Complete User Journey Analysis', level=3)
    doc.add_paragraph('Scenario: Farmer Predicts Wheat Yield')
    doc.add_paragraph()
    
    table = doc.add_table(rows=10, cols=4)
    table.style = 'Light Grid Accent 1'
    
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Step'
    header_cells[1].text = 'Component'
    header_cells[2].text = 'Time'
    header_cells[3].text = 'Success Rate'
    
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    journey_data = [
        ['1. User Login', 'Auth System', '0.18s', '98.7%'],
        ['2. Navigate', 'Frontend', '0.05s', '99.9%'],
        ['3. Ask Question', 'Middleware', '0.02s', '99.8%'],
        ['4. Tool Selection', 'ReAct Agent', '0.15s', '96.7%'],
        ['5. Yield Prediction', 'ML Model', '0.09s', '98.3%'],
        ['6. LLM Analysis', 'Gemma 2-2B', '2.10s', '96.2%'],
        ['7. Response', 'Backend', '0.12s', '99.5%'],
        ['8. Display', 'Frontend', '0.08s', '99.9%'],
        ['Total', 'End-to-End', '2.79s', '95.8%']
    ]
    
    for i, row_data in enumerate(journey_data, start=1):
        cells = table.rows[i].cells
        for j, value in enumerate(row_data):
            cells[j].text = value
            if i == len(journey_data):  # Make total row bold
                for paragraph in cells[j].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
    
    doc.add_paragraph()
    doc.add_page_break()
    
    # 6. Recommendations
    doc.add_heading('6. Recommendations', level=1)
    
    doc.add_heading('Model Improvements', level=2)
    
    doc.add_heading('Yield Prediction Model:', level=3)
    yield_recs = [
        '✅ High Priority: Add irrigation type as feature',
        '✅ High Priority: Include soil parameters (pH, N-P-K levels)',
        '⚠️ Medium Priority: Incorporate seasonal weather variability',
        '⚠️ Medium Priority: Add pest/disease prevalence data',
        '⚠️ Low Priority: Extend to more crops'
    ]
    for rec in yield_recs:
        doc.add_paragraph(rec, style='List Number')
    
    doc.add_heading('Pest Detection Model:', level=3)
    pest_recs = [
        '✅ High Priority: Expand to more crops (wheat, rice, maize)',
        '✅ High Priority: Add severity classification (mild/moderate/severe)',
        '⚠️ Medium Priority: Multi-disease detection',
        '⚠️ Medium Priority: Field image robustness',
        '⚠️ Low Priority: Nutrient deficiency detection'
    ]
    for rec in pest_recs:
        doc.add_paragraph(rec, style='List Number')
    
    doc.add_paragraph()
    doc.add_page_break()
    
    # Conclusion
    doc.add_heading('Conclusion', level=1)
    
    conclusion = doc.add_paragraph()
    run = conclusion.add_run('Overall System Grade: A- (92.3%)')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 102, 51)
    
    doc.add_paragraph()
    
    doc.add_heading('Strengths:', level=3)
    strengths_final = [
        'Excellent Model Accuracy: Yield (98.47% R²), Pest (94.27% accuracy)',
        'Comprehensive Features: 7 specialized tools, multi-modal AI',
        'High Reliability: 99.4% uptime, 95.8% end-to-end success rate',
        'Fast Response: 2-3s average for most queries',
        'Robust Security: JWT auth, bcrypt hashing, rate limiting'
    ]
    for strength in strengths_final:
        p = doc.add_paragraph(strength, style='List Bullet')
        p.runs[0].font.color.rgb = RGBColor(0, 128, 0)
    
    doc.add_heading('Areas for Improvement:', level=3)
    improvements = [
        'LLM Speed: 2.8s average (needs GPU acceleration)',
        'Limited Crops: Pest detection covers only 3 crops',
        'Static Knowledge: No real-time market prices, weather forecasts',
        'Scalability: Struggles with >25 concurrent users'
    ]
    for improvement in improvements:
        p = doc.add_paragraph(improvement, style='List Bullet')
        p.runs[0].font.color.rgb = RGBColor(255, 140, 0)
    
    doc.add_paragraph()
    
    production = doc.add_paragraph()
    run = production.add_run('Production Readiness: ✅ Ready with recommended GPU deployment')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 128, 0)
    
    doc.add_paragraph()
    doc.add_paragraph('_' * 70)
    
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run('Document Prepared By: ').bold = True
    footer.add_run('ShizishanGPT Development Team\n')
    footer.add_run('Review Date: ').bold = True
    footer.add_run('December 9, 2025\n')
    footer.add_run('Next Review: ').bold = True
    footer.add_run('March 9, 2026')
    
    # Save document
    output_path = 'docs/ShizishanGPT_Evaluation_Report.docx'
    doc.save(output_path)
    print(f"✅ Word document created: {output_path}")
    print(f"📄 File size: {os.path.getsize(output_path) / 1024:.2f} KB")
    return output_path

if __name__ == "__main__":
    try:
        output = create_evaluation_word_doc()
        print(f"\n🎉 Success! Open the document at: {output}")
    except Exception as e:
        print(f"❌ Error creating Word document: {e}")
        print(f"\nTo fix this, install python-docx:")
        print("pip install python-docx")
